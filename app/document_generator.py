"""Generate DOCX or PDF from title + content (e.g. text from OpenClaw research)."""

import io
from typing import Literal

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def generate_docx(title: str, content: str) -> bytes:
    """Build a DOCX with title and body; return bytes."""
    doc = Document()
    doc.add_heading(title, level=0)
    for para in content.strip().split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(8)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _escape_paragraph(s: str) -> str:
    """Escape <, >, & for ReportLab Paragraph (XML-style)."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_pdf(title: str, content: str) -> bytes:
    """Build a PDF with title and body; return bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(_escape_paragraph(title), styles["Title"]))
    story.append(Spacer(1, 12))
    for para in content.strip().split("\n\n"):
        if para.strip():
            safe = _escape_paragraph(para.strip()).replace("\n", "<br/>")
            story.append(Paragraph(safe, styles["Normal"]))
            story.append(Spacer(1, 6))
    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_document(title: str, content: str, fmt: Literal["docx", "pdf"]) -> tuple[bytes, str, str]:
    """
    Generate document bytes and suggested filename/content-type.
    Returns (bytes, filename, media_type).
    """
    if fmt == "docx":
        data = generate_docx(title, content)
        filename = _safe_filename(title) + ".docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "pdf":
        data = generate_pdf(title, content)
        filename = _safe_filename(title) + ".pdf"
        media_type = "application/pdf"
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return data, filename, media_type


def _safe_filename(title: str) -> str:
    """Turn title into a safe filename (no path, no special chars)."""
    safe = "".join(c if c.isalnum() or c in " -_" else "_" for c in title.strip())
    return (safe[:80] or "document").strip()
