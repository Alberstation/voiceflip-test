"""
Document loaders: DOCX, HTML, and PDF.
Return list of (content, metadata) for downstream chunking.
"""
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.cleaning import clean_text
from app.constants import SUPPORTED_DOC_EXTENSIONS


def _metadata_from_path(file_path: Path, source_id: str) -> dict[str, Any]:
    return {
        "source": str(file_path),
        "doc_id": source_id,
        "filename": file_path.name,
    }


def load_docx(path: Path, doc_id: str | None = None) -> list[tuple[str, dict[str, Any]]]:
    """
    Load a DOCX file. Returns one block per paragraph (optional to merge later).
    doc_id: optional identifier (default: stem of path).
    """
    doc_id = doc_id or path.stem
    doc = DocxDocument(path)
    parts: list[tuple[str, dict[str, Any]]] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        text = clean_text(text)
        if not text:
            continue
        meta = _metadata_from_path(path, doc_id) | {"page_or_para": i + 1}
        parts.append((text, meta))
    if not parts:
        full = "\n".join(p.text for p in doc.paragraphs)
        full = clean_text(full)
        if full:
            parts.append((full, _metadata_from_path(path, doc_id) | {"page_or_para": 1}))
    return parts


def load_pdf(path: Path, doc_id: str | None = None) -> list[tuple[str, dict[str, Any]]]:
    """
    Load a PDF file. One block per page (text extracted). doc_id: optional identifier.
    """
    doc_id = doc_id or path.stem
    reader = PdfReader(path)
    parts: list[tuple[str, dict[str, Any]]] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if not text or not text.strip():
            continue
        text = clean_text(text.strip())
        if not text:
            continue
        meta = _metadata_from_path(path, doc_id) | {"page_or_para": i + 1}
        parts.append((text, meta))
    return parts


def load_html(path: Path, doc_id: str | None = None) -> list[tuple[str, dict[str, Any]]]:
    """
    Load an HTML file. Extracts text from body; optional per-element blocks.
    doc_id: optional identifier (default: stem of path).
    """
    doc_id = doc_id or path.stem
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    body = soup.find("body") or soup
    text = body.get_text(separator="\n", strip=True)
    text = clean_text(text)
    if not text:
        return []
    meta = _metadata_from_path(path, doc_id) | {"page_or_para": 1}
    return [(text, meta)]


def load_document(path: Path, doc_id: str | None = None) -> list[tuple[str, dict[str, Any]]]:
    """Dispatch by extension: .docx -> load_docx, .pdf -> load_pdf, .html/.htm -> load_html."""
    suf = path.suffix.lower()
    if suf not in SUPPORTED_DOC_EXTENSIONS:
        raise ValueError(f"Unsupported format: {suf}. Use DOCX, PDF, or HTML.")
    if suf == ".pdf":
        return load_pdf(path, doc_id)
    if suf == ".docx":
        return load_docx(path, doc_id)
    return load_html(path, doc_id)
